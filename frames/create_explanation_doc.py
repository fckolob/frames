from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('calculateFrameBarsQuantityWithCustomLength Method Explanation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add overview section
doc.add_heading('Overview', 1)
p = doc.add_paragraph()
p.add_run('This method solves the ')
run = p.add_run('bin packing problem')
run.bold = True
p.add_run(': given a list of pieces to cut and a bar length, find the minimum number of bars needed. It uses a hybrid approach combining greedy and optimal algorithms.')

# Add main content
doc.add_heading('Step-by-Step Breakdown', 1)

# Step 1
doc.add_heading('Step 1: Initialize Variables (Lines 1143-1144)', 2)
code = doc.add_paragraph('const slice = 4;', style='Intense Quote')
doc.add_paragraph('• Sets a 4mm slice (saw blade kerf) that\'s added to each piece when cutting')

# Step 2
doc.add_heading('Step 2: Sort Pieces in Descending Order (Line 1147)', 2)
code = doc.add_paragraph('const pieces = [...lenghtGroup].sort((a, b) => b - a);', style='Intense Quote')
doc.add_paragraph('• Creates a copy of the input array and sorts pieces from largest to smallest')
doc.add_paragraph('• Sorting helps the algorithm work more efficiently (larger pieces first)')

# Step 3
doc.add_heading('Step 3: Get Initial Greedy Solution (Line 1150)', 2)
code = doc.add_paragraph('let bestSolution = greedyBinPacking(pieces, barLength, slice);', style='Intense Quote')
doc.add_paragraph('• Runs a fast greedy algorithm to get an initial solution')
doc.add_paragraph('• This serves as an upper bound for the optimal algorithm')
doc.add_paragraph('• The greedy approach: fit pieces into bins using a "first-fit" or similar heuristic')

# Step 4
doc.add_heading('Step 4: Threshold Check (Lines 1153-1155)', 2)
code = doc.add_paragraph('''if (pieces.length > 40) {
    return { quantity: bestSolution, method: "Greedy" };
}''', style='Intense Quote')
doc.add_paragraph('• Performance optimization: If there are more than 40 pieces, skip the optimal algorithm')
doc.add_paragraph('• Returns the greedy solution immediately (optimal would be too slow)')
doc.add_paragraph('• Returns an object with the quantity and method used')

# Step 5
doc.add_heading('Step 5: Setup for Branch and Bound Algorithm (Lines 1158-1159)', 2)
code = doc.add_paragraph('''const countRef = pieces.length;
const bins = new Float64Array(pieces.length);''', style='Intense Quote')
doc.add_paragraph('• countRef: Total number of pieces to pack')
doc.add_paragraph('• bins: Array to track remaining free space in each bin (bar)')
doc.add_paragraph('• Uses Float64Array for performance')

# Step 6
doc.add_heading('Step 6: Define the DFS Branch-and-Bound Function (Lines 1161-1220)', 2)
doc.add_paragraph('This is the core optimal algorithm. The dfs_bnb function contains several key components:')

# Step 6a
doc.add_heading('6a. Pruning 1 - Early Exit (Line 1162)', 3)
code = doc.add_paragraph('if (binCount >= bestSolution) return;', style='Intense Quote')
doc.add_paragraph('• If current bin count already equals or exceeds the best solution found, stop exploring this branch')
doc.add_paragraph('• No point continuing if we can\'t improve')

# Step 6b
doc.add_heading('6b. Base Case - All Pieces Placed (Lines 1165-1170)', 3)
code = doc.add_paragraph('''if (currentPieceIdx >= countRef) {
    if (binCount < bestSolution) {
        bestSolution = binCount;
    }
    return;
}''', style='Intense Quote')
doc.add_paragraph('• If we\'ve placed all pieces (currentPieceIdx reached the end)')
doc.add_paragraph('• Update bestSolution if this configuration uses fewer bins')
doc.add_paragraph('• Return to explore other branches')

# Step 6c
doc.add_heading('6c. Pruning 2 - Lower Bound Calculation (Lines 1173-1191)', 3)
code = doc.add_paragraph('''let remainingSum = 0;
for (let i = currentPieceIdx; i < countRef; i++) {
    remainingSum += (pieces[i] + slice);
}

let currentFreeSpace = 0;
for (let i = 0; i < binCount; i++) {
    currentFreeSpace += bins[i]; 
}

let neededVolume = remainingSum - currentFreeSpace;
let minAdditional = 0;
if (neededVolume > 0) {
    minAdditional = Math.ceil(neededVolume / barLength);
}

if (binCount + minAdditional >= bestSolution) return;''', style='Intense Quote')
doc.add_paragraph('• Calculates a lower bound on how many additional bins we\'ll need')
doc.add_paragraph('• remainingSum: Total length of all remaining pieces (including slice)')
doc.add_paragraph('• currentFreeSpace: Total free space in currently used bins')
doc.add_paragraph('• neededVolume: How much more space we need')
doc.add_paragraph('• minAdditional: Minimum additional bins required (theoretical minimum)')
doc.add_paragraph('• If current bins + minimum needed >= bestSolution, prune this branch')

# Step 6d
doc.add_heading('6d. Get Current Piece Size (Line 1193)', 3)
code = doc.add_paragraph('const pieceSize = pieces[currentPieceIdx] + slice;', style='Intense Quote')
doc.add_paragraph('• Add the slice (kerf) to the piece length')

# Step 6e
doc.add_heading('6e. Try Placing in Existing Bins (Lines 1196-1211)', 3)
code = doc.add_paragraph('''for (let i = 0; i < binCount; i++) {
    // Pruning 3: Symmetry Breaking
    let symmetric = false;
    for(let k=0; k < i; k++){
        if(Math.abs(bins[k] - bins[i]) < 0.001){
            symmetric = true;
            break;
        }
    }
    if(symmetric) continue;

    if (bins[i] >= pieceSize) {
        bins[i] -= pieceSize;
        dfs_bnb(currentPieceIdx + 1, binCount);
        bins[i] += pieceSize; // Backtrack
        
        if (bestSolution <= binCount) return;
    }
}''', style='Intense Quote')
doc.add_paragraph('• Symmetry Breaking: Skip bins with identical free space (they\'re equivalent choices)')
doc.add_paragraph('• For each bin with enough space:')
p = doc.add_paragraph('  - Place the piece (reduce free space)', style='List Bullet 2')
p = doc.add_paragraph('  - Recursively try to place the next piece', style='List Bullet 2')
p = doc.add_paragraph('  - Backtrack (restore free space) to try other options', style='List Bullet 2')
p = doc.add_paragraph('  - Early exit if we\'ve found a better solution', style='List Bullet 2')

# Step 6f
doc.add_heading('6f. Try Creating a New Bin (Lines 1214-1217)', 3)
code = doc.add_paragraph('''if (binCount + 1 < bestSolution) {
    bins[binCount] = barLength - pieceSize;
    dfs_bnb(currentPieceIdx + 1, binCount + 1);
}''', style='Intense Quote')
doc.add_paragraph('• Only create a new bin if it could potentially improve the solution')
doc.add_paragraph('• Set the new bin\'s free space to barLength - pieceSize')
doc.add_paragraph('• Recursively continue with the next piece')

# Step 7
doc.add_heading('Step 7: Execute the Algorithm and Return (Lines 1220-1222)', 2)
code = doc.add_paragraph('''dfs_bnb(0, 0);
return { quantity: bestSolution, method: "Optimal" };''', style='Intense Quote')
doc.add_paragraph('• Start the recursive algorithm with piece 0 and 0 bins')
doc.add_paragraph('• Return the optimal solution with the method label "Optimal"')

# Summary
doc.add_heading('Summary', 1)
doc.add_paragraph('This method uses a two-phase approach:')
doc.add_paragraph('Greedy phase: Quick approximation for all cases', style='List Number')
doc.add_paragraph('Optimal phase: Branch-and-bound DFS for ≤40 pieces', style='List Number')

doc.add_paragraph()
doc.add_paragraph('The optimal algorithm uses three pruning techniques:')
doc.add_paragraph('Pruning 1: Stop if current solution can\'t beat the best', style='List Bullet')
doc.add_paragraph('Pruning 2: Calculate lower bounds to eliminate impossible branches', style='List Bullet')
doc.add_paragraph('Pruning 3: Symmetry breaking to avoid redundant exploration', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('The result is an efficient bin packing solver that balances ')
p.add_run('speed').bold = True
p.add_run(' and ')
p.add_run('optimality').bold = True
p.add_run('! 🎯')

# Save the document
output_path = r'c:\Users\usuario\Documents\frames\frames\frames\calculateFrameBarsQuantityWithCustomLength_Explanation.docx'
doc.save(output_path)
print(f"Document created successfully at: {output_path}")
